"""
Document Content Extractors - Support for multiple document formats
Secure content extraction from PDF, DOC, DOCX, TXT, MD, and HTML files
"""

import logging
import io
import tempfile
import os
from typing import Optional, Dict, Any, Tuple
from pathlib import Path

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """
    Base class for document content extraction.
    Provides secure content extraction with format validation.
    """
    
    def __init__(self):
        """Initialize document extractor with security settings"""
        self.max_file_size = 50 * 1024 * 1024  # 50MB limit
        self.supported_formats = {'pdf', 'doc', 'docx', 'txt', 'md', 'html'}
    
    async def extract_content(self, file_data: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text content from document bytes.
        
        Args:
            file_data: Raw file bytes
            filename: Original filename for format detection
            
        Returns:
            Tuple of (extracted_text, metadata)
            
        Raises:
            ValueError: If file format unsupported or extraction fails
            RuntimeError: If file size exceeds limits
        """
        # Security validation
        if len(file_data) > self.max_file_size:
            raise RuntimeError(f"File size {len(file_data)} bytes exceeds limit {self.max_file_size}")
        
        if len(file_data) == 0:
            raise ValueError("Empty file provided")
        
        # Detect file format from extension
        file_extension = self._get_file_extension(filename)
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Extract content based on format
        try:
            if file_extension == 'pdf':
                return await self._extract_pdf(file_data)
            elif file_extension in ['doc', 'docx']:
                return await self._extract_word(file_data, file_extension)
            elif file_extension == 'txt':
                return await self._extract_text(file_data)
            elif file_extension == 'md':
                return await self._extract_markdown(file_data)
            elif file_extension == 'html':
                return await self._extract_html(file_data)
            else:
                raise ValueError(f"No extractor available for format: {file_extension}")
                
        except Exception as e:
            logger.error(f"Content extraction failed for {filename}: {e}")
            raise ValueError(f"Failed to extract content from {file_extension} file: {str(e)}")
    
    def _get_file_extension(self, filename: str) -> str:
        """Extract and validate file extension"""
        extension = Path(filename).suffix.lower().lstrip('.')
        return extension if extension else ''
    
    async def _extract_pdf(self, file_data: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract content from PDF file using PyPDF2"""
        try:
            import PyPDF2
        except ImportError:
            raise RuntimeError("PyPDF2 not available for PDF extraction")
        
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_data))
            
            # Security check for password-protected PDFs
            if pdf_reader.is_encrypted:
                raise ValueError("Password-protected PDFs are not supported")
            
            text_content = []
            page_count = len(pdf_reader.pages)
            
            # Limit number of pages for security
            max_pages = 1000
            if page_count > max_pages:
                logger.warning(f"PDF has {page_count} pages, limiting to {max_pages}")
                page_count = max_pages
            
            for page_num in range(page_count):
                try:
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text)
                except Exception as e:
                    logger.warning(f"Failed to extract page {page_num}: {e}")
                    continue
            
            extracted_text = '\n\n'.join(text_content)
            
            metadata = {
                'format': 'pdf',
                'pages': page_count,
                'extraction_method': 'PyPDF2'
            }
            
            return extracted_text, metadata
            
        except Exception as e:
            raise ValueError(f"PDF extraction failed: {str(e)}")
    
    async def _extract_word(self, file_data: bytes, format_type: str) -> Tuple[str, Dict[str, Any]]:
        """Extract content from Word documents using python-docx"""
        try:
            import docx
        except ImportError:
            raise RuntimeError("python-docx not available for Word document extraction")
        
        try:
            # Only support DOCX format for security reasons
            if format_type == 'doc':
                raise ValueError("Legacy DOC format not supported, please convert to DOCX")
            
            doc = docx.Document(io.BytesIO(file_data))
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            extracted_text = '\n\n'.join(text_content)
            
            metadata = {
                'format': 'docx',
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'extraction_method': 'python-docx'
            }
            
            return extracted_text, metadata
            
        except Exception as e:
            raise ValueError(f"Word document extraction failed: {str(e)}")
    
    async def _extract_text(self, file_data: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract content from plain text file"""
        try:
            # Try UTF-8 first, then common encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin1', 'cp1252']
            
            for encoding in encodings:
                try:
                    extracted_text = file_data.decode(encoding)
                    metadata = {
                        'format': 'txt',
                        'encoding': encoding,
                        'extraction_method': 'direct'
                    }
                    return extracted_text, metadata
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Could not decode text file with any supported encoding")
            
        except Exception as e:
            raise ValueError(f"Text extraction failed: {str(e)}")
    
    async def _extract_markdown(self, file_data: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract content from Markdown file"""
        try:
            # Markdown is plain text, use same extraction as TXT
            extracted_text, _ = await self._extract_text(file_data)
            
            metadata = {
                'format': 'md',
                'extraction_method': 'markdown'
            }
            
            return extracted_text, metadata
            
        except Exception as e:
            raise ValueError(f"Markdown extraction failed: {str(e)}")
    
    async def _extract_html(self, file_data: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract content from HTML file using BeautifulSoup"""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise RuntimeError("BeautifulSoup4 not available for HTML extraction")
        
        try:
            # Decode HTML content
            html_content, _ = await self._extract_text(file_data)
            
            # Parse HTML and extract text
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for element in soup(['script', 'style', 'meta', 'link']):
                element.decompose()
            
            # Extract text content
            extracted_text = soup.get_text(separator='\n', strip=True)
            
            # Clean up excessive whitespace
            lines = [line.strip() for line in extracted_text.split('\n') if line.strip()]
            extracted_text = '\n'.join(lines)
            
            metadata = {
                'format': 'html',
                'title': soup.title.string if soup.title else None,
                'extraction_method': 'BeautifulSoup'
            }
            
            return extracted_text, metadata
            
        except Exception as e:
            raise ValueError(f"HTML extraction failed: {str(e)}")


async def create_document_extractor() -> DocumentExtractor:
    """
    Factory function to create document extractor instance.
    
    Returns:
        Configured DocumentExtractor instance
    """
    return DocumentExtractor()